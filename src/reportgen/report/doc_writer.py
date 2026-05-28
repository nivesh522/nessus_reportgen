from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor
from loguru import logger

from ..models import Finding


def cm_to_twips(cm: float) -> int:
    return int(cm * 567)  # 1 cm = 567 twips


def set_cell_shading(cell, fill_color: str) -> None:
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(
        r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="%s"/>'
        % fill_color
    )
    tcPr.append(shd)


def set_cell_width(cell, width: int, width_type: str = "dxa") -> None:
    tcPr = cell._element.get_or_add_tcPr()
    tcW = parse_xml(
        r'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="%d" w:type="%s"/>'
        % (width, width_type)
    )
    tcPr.append(tcW)


def set_row_height(row, height_twips: int, rule: str = "auto") -> None:
    trPr = row._element.get_or_add_trPr()
    trHeight = parse_xml(
        r'<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="%d" w:hRule="%s"/>'
        % (height_twips, rule)
    )
    trPr.append(trHeight)


def enable_text_wrap(cell) -> None:
    tcPr = cell._element.get_or_add_tcPr()
    text_wrap = parse_xml(
        r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="0"/>'
    )
    tcPr.append(text_wrap)


SEVERITY_COLORS = {
    "Critical": "C0504D",
    "High": "F79646",
    "Medium": "F0F04E",
    "Low": "9BBB59",
    "Info": "4BACC6",
    "Informational": "4BACC6",
}
HEADER_FILL_COLOR = "73A0DD"
VULN_RESULT_COLUMNS = [cm_to_twips(3.69), cm_to_twips(19.55)]
OPEN_PORTS_COLUMNS = [cm_to_twips(3.19), cm_to_twips(10.02), cm_to_twips(10.02)]
REPORT_CARD_COL_WIDTH = cm_to_twips(5.81)
REPORT_CARD_ROW1_HEIGHT = cm_to_twips(2.29)
OPEN_PORTS_ROW_HEIGHT = cm_to_twips(0.47)
VULN_RESULT_ROW_HEIGHT = cm_to_twips(0.45)


class WordReportWriter:
    def __init__(self, template_path: Path | None = None):
        self.template_path = template_path

    def write(self, output_path: Path, findings: list[Finding]) -> None:
        if self.template_path and self.template_path.exists():
            doc = Document(str(self.template_path))
        else:
            doc = Document()

        self._add_vulnerability_report_card(doc, findings)
        self._add_host_summary(doc, findings)
        self._add_open_ports(doc, findings)
        self._add_findings(doc, findings)

        doc.save(str(output_path))
        logger.info(f"Word report written to {output_path}")

    def _add_vulnerability_report_card(self, doc: Document, findings: list[Finding]) -> None:
        severity_counts: dict[str, int] = {
            "Critical": 0,
            "High": 0,
            "Medium": 0,
            "Low": 0,
        }
        for f in findings:
            sev = f.severity.capitalize()
            if sev in severity_counts:
                severity_counts[sev] += 1

        doc.add_heading("Vulnerability Report Card", level=2)
        table = doc.add_table(rows=2, cols=4)
        set_row_height(table.rows[0], REPORT_CARD_ROW1_HEIGHT, "exact")

        for row in table.rows:
            for cell in row.cells:
                set_cell_width(cell, REPORT_CARD_COL_WIDTH)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.bold = True

        header_cells = table.rows[1].cells
        header_cells[0].text = "Critical"
        header_cells[1].text = "High"
        header_cells[2].text = "Medium"
        header_cells[3].text = "Low"
        set_cell_shading(header_cells[0], "000000")
        set_cell_shading(header_cells[1], "000000")
        set_cell_shading(header_cells[2], "000000")
        set_cell_shading(header_cells[3], "000000")
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)

        data_cells = table.rows[0].cells
        data_cells[0].text = str(severity_counts["Critical"])
        data_cells[1].text = str(severity_counts["High"])
        data_cells[2].text = str(severity_counts["Medium"])
        data_cells[3].text = str(severity_counts["Low"])
        set_cell_shading(data_cells[0], SEVERITY_COLORS["Critical"])
        set_cell_shading(data_cells[1], SEVERITY_COLORS["High"])
        set_cell_shading(data_cells[2], SEVERITY_COLORS["Medium"])
        set_cell_shading(data_cells[3], SEVERITY_COLORS["Low"])
        for cell in data_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)

        doc.add_paragraph()

    def _add_host_summary(self, doc: Document, findings: list[Finding]) -> None:
        host_data: dict[str, dict[str, int]] = {}
        all_ips: list[str] = []
        for f in findings:
            for ip in f.affected_ips:
                if ip not in host_data:
                    host_data[ip] = {
                        "Total": 0,
                        "Critical": 0,
                        "High": 0,
                        "Medium": 0,
                        "Low": 0,
                        "Info": 0,
                    }
                    all_ips.append(ip)
                host_data[ip]["Total"] += 1
                sev = f.severity.capitalize()
                if sev in host_data[ip]:
                    host_data[ip][sev] += 1
                else:
                    host_data[ip]["Info"] += 1

        doc.add_heading("Identified Hosts", level=2)
        doc.add_paragraph(
            f"In total {len(all_ips)} hosts were identified during the scan that had vulnerabilities in them."
        )
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "IP"
        for ip in sorted(all_ips):
            row_cells = table.add_row().cells
            row_cells[0].text = ip
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
        doc.add_paragraph()

        doc.add_heading("Vulnerabilities", level=2)
        doc.add_heading("Vulnerabilities and Affected Hosts", level=3)

    def _add_open_ports(self, doc: Document, findings: list[Finding]) -> None:
        host_ports: dict[str, list[str]] = {}
        all_ips = set()
        for f in findings:
            for ip in f.affected_ips:
                all_ips.add(ip)
                if ip not in host_ports:
                    host_ports[ip] = []
                host_ports[ip].extend(f.ports)
        for ip in host_ports:
            host_ports[ip] = sorted(list(set(host_ports[ip])))

        doc.add_heading("Open Ports", level=2)
        table = doc.add_table(rows=1, cols=3)
        header_cells = table.rows[0].cells
        header_cells[0].text = "Host IP"
        header_cells[1].text = "Open TCP Ports"
        header_cells[2].text = "Open UDP Ports"

        for i, cell in enumerate(header_cells):
            set_cell_width(cell, OPEN_PORTS_COLUMNS[i])
            set_cell_shading(cell, HEADER_FILL_COLOR)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

        for ip in sorted(all_ips):
            row = table.add_row()
            set_row_height(row, OPEN_PORTS_ROW_HEIGHT, "auto")
            row.cells[0].text = ip
            row.cells[1].text = ", ".join(host_ports.get(ip, []))
            row.cells[2].text = "-"
            for i, cell in enumerate(row.cells):
                set_cell_width(cell, OPEN_PORTS_COLUMNS[i])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

        doc.add_paragraph()

    def _add_findings(self, doc: Document, findings: list[Finding]) -> None:
        for finding in findings:
            if finding.severity.lower() == "info":
                continue
            table = doc.add_table(rows=8, cols=2)
            fields = [
                ("Synopsis", finding.synopsis or "N/A"),
                ("Ports", ", ".join(finding.ports) if finding.ports else "N/A"),
                ("Severity", finding.severity),
                ("Risk", finding.risk or "N/A"),
                ("Description", finding.description or "N/A"),
                ("Solution", finding.solution or "N/A"),
                ("CVE numbers", "\n".join(finding.cves) if finding.cves else ""),
                (
                    "Affected IP addresses",
                    ", ".join(finding.affected_ips) if finding.affected_ips else "N/A",
                ),
            ]
            for i, (label, value) in enumerate(fields):
                cell_left = table.rows[i].cells[0]
                cell_right = table.rows[i].cells[1]
                cell_left.text = label
                cell_right.text = value

                set_row_height(table.rows[i], VULN_RESULT_ROW_HEIGHT, "auto")

                set_cell_width(cell_left, VULN_RESULT_COLUMNS[0])
                set_cell_width(cell_right, VULN_RESULT_COLUMNS[1])

                enable_text_wrap(cell_left)
                enable_text_wrap(cell_right)

                for paragraph in cell_left.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.bold = True
                        if i == 0:
                            run.font.color.rgb = RGBColor(255, 255, 255)

                for paragraph in cell_right.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

                if i == 0:
                    set_cell_shading(cell_left, HEADER_FILL_COLOR)
                    set_cell_shading(cell_right, HEADER_FILL_COLOR)

                if i in [2, 3]:
                    sev_key = finding.severity.capitalize()
                    if sev_key in SEVERITY_COLORS:
                        set_cell_shading(cell_right, SEVERITY_COLORS[sev_key])

            doc.add_paragraph()
